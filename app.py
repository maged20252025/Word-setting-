
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

TEMPLATE_PATH = "دعوى استحقاق.docx"

st.set_page_config(page_title="منسق الدعاوى بنموذج", layout="centered")
st.title("📄 منسق الدعاوى بنمط موحد")

uploaded_doc = st.file_uploader("📤 ارفع ملف الدعوى (Word غير منسق)", type=["docx"])

if uploaded_doc and st.button("🔧 تنسيق الملف تلقائيًا"):
    try:
        # قراءة النص من الملف غير المنسق
        input_doc = Document(uploaded_doc)
        text_lines = [p.text for p in input_doc.paragraphs if p.text.strip()]

        # تحميل القالب واستخراج النمط من أول فقرة
        template_doc = Document(TEMPLATE_PATH)
        reference_style = template_doc.paragraphs[0].style

        # حذف جميع الفقرات الموجودة في القالب
        while len(template_doc.paragraphs):
            p = template_doc.paragraphs[0]
            p._element.getparent().remove(p._element)

        # إدراج النص الجديد بنفس تنسيق الفقرة الأصلية
        for line in text_lines:
            para = template_doc.add_paragraph(style=reference_style)
            run = para.add_run(line.strip())
            run.font.name = 'Traditional Arabic'
            run.font.size = Pt(16)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # حفظ الملف داخل BytesIO
        output = io.BytesIO()
        template_doc.save(output)
        output.seek(0)

        st.success("✅ تم تنسيق الملف بنجاح!")
        st.download_button("📥 تحميل الملف المنسق", data=output, file_name="دعوى_منسقة.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        st.error("❌ حدث خطأ أثناء التنسيق:")
        st.exception(e)
