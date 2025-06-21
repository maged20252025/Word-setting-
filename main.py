import streamlit as st
from docx import Document
from io import BytesIO
import os

st.set_page_config(page_title="منسق المستندات القانونية", layout="centered")

st.title("📄 منسق المستندات القانونية")
st.markdown("قم برفع ملف Word غير منسق واختر قالبًا منسقًا ليتم تطبيق تنسيقه عليه تلقائيًا.")

# رفع الملف غير المنسق
uploaded_file = st.file_uploader("📤 ارفع الملف غير منسق (.docx)", type=["docx"])

# اختيار القالب
TEMPLATE_DIR = "templates"
template_files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx")]
selected_template = st.selectbox("📑 اختر القالب:", template_files)

# عند الضغط على زر التنسيق
if uploaded_file and selected_template:
    if st.button("✅ تنسيق الملف"):
        try:
            # تحميل الملف غير المنسق
            original_doc = Document(uploaded_file)

            # تحميل القالب
            template_path = os.path.join(TEMPLATE_DIR, selected_template)
            template_doc = Document(template_path)

            # استخراج أول فقرة كنموذج للتنسيق
            if not template_doc.paragraphs:
                st.error("⚠️ القالب المحدد لا يحتوي على فقرات!")
                st.stop()

            sample_para = template_doc.paragraphs[0]
            style = sample_para.style

            # إنشاء مستند جديد بنفس تنسيق الفقرة الأولى من القالب
            new_doc = Document()
            for para in original_doc.paragraphs:
                if para.text.strip():
                    new_para = new_doc.add_paragraph(para.text, style=style)

            # حفظ النتيجة
            output = BytesIO()
            new_doc.save(output)
            output.seek(0)

            st.success("✅ تم تنسيق الملف بنجاح!")
            st.download_button("📥 تحميل الملف المنسق", output, file_name="ملف_منسق.docx")
        except Exception as e:
            st.error(f"حدث خطأ أثناء التنسيق: {str(e)}")
